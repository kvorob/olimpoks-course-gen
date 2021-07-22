import docx


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText


def check_answer(str):
    i = 0
    while i < len(str):
        if str[i] == '+': return 1
        if str[i].isalpha(): return 0
        i += 1


def get_help(str):
    i = 0
    buf = ''
    hlp = ''
    while i < len(str):
        if str[i] != '.':
            buf += str[i]
        else:
            if i < (len(str) // 2):
                hlp += buf + '.'
                buf = ''
        i += 1
    return [hlp.strip(), buf.strip()]


def parse_cell_by_bold_attr(qcell, count, qs):
    if qcell._tc.grid_span > 1: return
    i = 0
    q = []
    a = []
    buf = ''
    quest = ''
    for pr in qcell.paragraphs:
        if pr.style.name.find('Heading') != -1:
            buf += pr.text
        else:
            for run in pr.runs:
                if run.bold: buf += run.text
        if len(buf) > 0 and i == 0:
            quest += buf + " "
            buf = ''
            continue
        if len(buf) == 0 and i == 0:  # Вопрос закончился
            hlp = get_help(quest)
            q = [len(qs)+1] + hlp
            i += 1
            a.append([i, check_answer(pr.text), pr.text.strip(' +')])
        else:
            a.append([i, check_answer(pr.text), pr.text.strip(' +')])
            i += 1
    if q:
        q.append(a)
        qs.append(q)
    return


def tables_info(table_item, col_id, qs):
    count = 0
    for row_idx in range(len(table_item.rows)):
        for col_idx in range(len(table_item.columns)):
            if col_idx == col_id:
                parse_cell_by_bold_attr(table_item.cell(row_idx, col_idx), count,qs)
                count += 1
    return

def show_tables(filename):
    doc = docx.Document(filename)
    for t in doc.tables:
        if len(t.rows) < 10: continue  # Таблицы в которых меньше 20 строк не обрабатываем
        for row_idx in range(len(t.rows)):
            for col_idx in range(len(t.columns)):
                print('Row: {} Col:{} Content: {}'.format(row_idx,col_idx, t.cell(row_idx, col_idx).text))
								
    return


def parse_table(filename, col_id, qs):
    doc = docx.Document(filename)
    for t in doc.tables:
        if len(t.rows) < 10: continue  # Таблицы в которых меньше 20 строк не обрабатываем
        tables_info(t, col_id, qs)
    return

if __name__ == '__main__':  
   questions = []
   col_id = 3
   filename = 'src//043-18.docx'
   show_tables(filename)
   parse_table(filename, col_id, questions)
   print(questions)

