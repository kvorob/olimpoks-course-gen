import docx
from logger import TLogger
from lxml import etree
from os.path import basename, splitext
import binascii, base64


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText


def check_answer(str):
    i = 0
    while i < len(str):
        if str[i] == '+': return 1
        if str[i].isalpha(): return 0
        i += 1


def get_question_and_help_list(str,hcell):
    i = 0
    buf = ''
    hlp = ''
    while i < len(str):
        if str[i] != '.':
            buf += str[i]
        else:
            if i < (len(str) // 2):
                hlp += buf + '.'
                buf = ''
        i += 1
    if hcell:
        #Пришла ячейка с текстом помощи. Читаем ее
        hbuf = ""
        for pr in hcell.paragraphs:
            hbuf = hbuf+'<p>'+pr.text+'</p>'
        print(f'Help:{hbuf}')
        hlp = hbuf
    return [hlp.strip(), buf.strip()]


def get_imag2(run, doc):
    #ToDo Сделать режим с преобразованием размера иллюстрации до рекомендованного
    imageURL = ""
    if 'graphicData' in run._r.xml:
        rt = etree.ElementTree(etree.XML(run._r.xml)).getroot()
        for item in rt.findall(".//{*}graphicData/{*}pic/{*}blipFill/{*}blip"):
            for k in item.keys():
                if "embed" in k: rId = item.get(k)
            try:
                contentType = doc.part.related_parts[rId].content_type
            except Exception as e:
                TLogger.WriteLog("Error", "Ошибка поиска идентификатора картинки в ячейке: {}".format(e))
                continue
            if not contentType.startswith("image"): continue
            imageName = basename(doc.part.related_parts[rId].partname)
            imageExt = imageName.split(".")[1]
            imageData = doc.part.related_parts[rId]._blob
            base64_bytes = base64.b64encode(imageData)
            base64_message = base64_bytes.decode('ascii')
            imageURL = f'<img src="data:image/{imageExt};base64,{base64_message}"/>'
    return imageURL


def parse_cell_by_bold_attr(qcell, hcell, count, qs, doc):
    if qcell._tc.grid_span > 1: return
    i = 0
    q = []
    a = []
    quest = ''
    q_images = []
    for pr in qcell.paragraphs:
        is_true = 0
        buf = ''
        q_buf = ''
        for run in pr.runs:
            # иллюстрации из вопроса упаковываем в QuestionMainImg, но только если в тексте вопроса всего одна илл.
            if i == 0:
                q_buf += run.text
                img = get_imag2(run, doc)
                buf += img + run.text
                if len(img) > 0: q_images.append(img)
            else:
                buf += get_imag2(run, doc) + run.text
            if run.bold: is_true = 1
        if len(buf) > 0 and i == 0:  # Пришел первый абзац, значит это вопрос
            if len(q_images) == 1:
                quest += q_buf + " "
            else:
                quest += buf + " "
            i += 1
            # Возвращаем сразу лист в котором первый элемент вопрос, второй помощь
            qh_lst = get_question_and_help_list(quest, hcell)
            # Если в вопросе есть только одна картинка упаковываем ее в дополнительное поле
            if len(q_images) == 0 or len(q_images) > 1: q_images = ["none"]
            q = [len(qs) + 1] + qh_lst +q_images
            #print(f"Img: {len(q_images)} Вопрос: {q}")
            continue
        elif len(buf) > 0:
            a.append([i, is_true, buf.strip()])
        i += 1
    if q:
        q.append(a)
        qs.append(q)
    return


def tables_info(table_item, col_id, qs, doc):
    count = 0
    for row_idx in range(len(table_item.rows)):
        for col_idx in range(len(table_item.columns)):
            if col_idx == col_id:
                if len(table_item.columns)>col_idx+1:
                    parse_cell_by_bold_attr(table_item.cell(row_idx, col_idx), table_item.cell(row_idx, col_idx+1), count, qs, doc)
                else:
                    parse_cell_by_bold_attr(table_item.cell(row_idx, col_idx),None, count, qs, doc)
                count += 1
    return


def show_tables(filename):
    doc = docx.Document(filename)
    for t in doc.tables:
        if len(t.rows) < 10: continue  # Таблицы в которых меньше 20 строк не обрабатываем
        for row_idx in range(len(t.rows)):
            for col_idx in range(len(t.columns)):
                print('Row: {} Col:{} Content: {}'.format(row_idx, col_idx, t.cell(row_idx, col_idx).text))

    return


def parse_table(filename, qs):
    doc = docx.Document(filename)
    for t in doc.tables:
        tables_info(t, 0, qs, doc)
    return


if __name__ == '__main__':
    questions = []
    col_id = 3
    filename = 'src//043-18.docx'
    show_tables(filename)
    parse_table(filename, col_id, questions)
    print(questions)
